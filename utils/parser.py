"""
Resume parsing utilities.
Extracts raw text from PDF, DOCX, or TXT resume files.
"""
import io
import docx
import pdfplumber


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file given as bytes."""
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file given as bytes."""
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also grab text inside tables (some resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file given as bytes."""
    return file_bytes.decode("utf-8", errors="ignore")


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    Raises ValueError for unsupported file types.
    """
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif lower_name.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: {filename}. Please upload a PDF, DOCX, or TXT file."
        )
